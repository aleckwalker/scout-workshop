"""Generate the workshop advertising doc from the cleaned-up markdown."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rocky Mountain Stanford Association")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Vibe-Coding Workshop")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("April 18, 2026  \u00b7  1:00 - 5:00 PM")
run.font.size = Pt(12)

# --- What's This About? ---
doc.add_heading("What\u2019s This About?", level=1)
doc.add_paragraph(
    "Building software used to be inaccessible to many, but the landscape is changing "
    "quickly. Vibe-coding is an emerging approach where you describe what you want in "
    "plain English and AI writes the code for you. It has lowered the floor enough that "
    "people without technical backgrounds can attend a workshop like this one and emerge "
    "capable of completing genuinely useful personal projects on their own. A session "
    "like this can take someone from curious observer to hobbyist builder, capable of "
    "making real things, without much more continued guidance beyond what we cover right here."
)
doc.add_paragraph(
    "This workshop is a hands-on introduction to that world. In four hours, you will go "
    "from zero to a working, personalized tool you can start using the same day."
)

# --- What Are We Building? ---
doc.add_heading("What Are We Building?", level=1)
doc.add_paragraph(
    "We thought long and hard about how to align the audience on a tool that creates "
    "value for everyone. We figure everyone has a niche interest or business topic that "
    "they wish they were more informed about from true experts in the field, and we\u2019re "
    "aware that there\u2019s more noise out there than ever. Trying to get up to speed on a "
    "new industry? Tracking a technology you want to understand better? Staying current "
    "in your own field without spending hours on it? Tell the tool your topic and it "
    "finds the experts, monitors what they publish, and brings the signal to you."
)

p = doc.add_paragraph()
run = p.add_run("The Goal (~3 hours, zero cost): ")
run.bold = True
p.add_run(
    "Build a tool that discovers the leading voices on any topic across Substack, Medium, "
    "Reddit, YouTube, and Twitter/X, lets you vet them with real evidence, monitors their "
    "RSS feeds, and generates weekly digests. You build it, you own it, you take it home. "
    "It costs nothing to run."
)

p = doc.add_paragraph()
run = p.add_run("The Stretch Goal (requires API key, ~$0.01/use): ")
run.bold = True
p.add_run(
    "Add AI-powered engagement suggestions. Your tool not only reads what thought leaders "
    "are posting but drafts platform-specific responses in your voice, turning passive "
    "monitoring into active participation."
)

p = doc.add_paragraph()
run = p.add_run("You\u2019ll leave with two tools:")
run.bold = True
doc.add_paragraph(
    "The one you build yourself, customized to your topic and goals.", style="List Bullet"
)
doc.add_paragraph(
    "The instructor\u2019s polished version (goal + stretch), so no matter what, you walk "
    "out with something useful.",
    style="List Bullet",
)

# --- Who Is This For? ---
doc.add_heading("Who Is This For?", level=1)
doc.add_paragraph(
    "You don\u2019t need to know how to code. If you have ever thought \u201cI wish there "
    "was a tool that did X,\u201d this workshop is for you."
)

audiences = [
    (
        "If you\u2019re non-technical:",
        " You\u2019ll leave with a working, personalized tool and a new mental model for "
        "what\u2019s now possible to build. No prior coding experience required.",
    ),
    (
        "If you\u2019re a founder or entrepreneur:",
        " Learn how to rapidly prototype tools that support your fundraising, client "
        "development, or growth strategy, without needing to hire a developer.",
    ),
    (
        "If you\u2019re in marketing, sales, or business development:",
        " Walk away with a live tool that finds the leading voices in your space across "
        "multiple platforms and keeps you current on what they are saying, so you always "
        "know where the conversation is headed.",
    ),
    (
        "If you have a technical background:",
        " This is a great opportunity to geek out on architecture, prompt engineering, "
        "and the emerging workflow of AI-assisted development, and compare notes with "
        "fellow builders.",
    ),
    (
        "If you\u2019re simply curious:",
        " About how AI is transforming software development and want a front-row seat to "
        "what that looks like in practice, this is your session.",
    ),
]
for bold_part, rest in audiences:
    p = doc.add_paragraph()
    run = p.add_run(bold_part)
    run.bold = True
    p.add_run(rest)

# --- Agenda ---
doc.add_heading("Agenda", level=1)
agenda = [
    ("1:00", "Welcome, setup check, pick your topic"),
    ("1:20", "Prompt 1: Project scaffolding & configuration"),
    ("1:40", "Prompt 2: Discovery engine \u2014 find voices across 5 platforms"),
    ("2:25", "Prompt 3: Approval flow \u2014 vet candidates with real evidence"),
    ("2:55", "Break"),
    ("3:05", "Prompt 4: RSS monitoring \u2014 follow your approved voices"),
    ("3:35", "Prompt 5: Digest generation \u2014 weekly summaries"),
    ("4:15", "Prompt 6: Customize, extend, or start on the stretch goal"),
    ("4:50", "Wrap-up: what you built, next steps, Q&A"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Shading Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Time"
hdr[1].text = "Section"
for time_str, section in agenda:
    row = table.add_row().cells
    row[0].text = time_str
    row[1].text = section

# --- Facilitators ---
doc.add_heading("Your Facilitators", level=1)

p = doc.add_paragraph()
run = p.add_run("Alec Walker")
run.bold = True
doc.add_paragraph(
    "Alec is a technology investor and serial entrepreneur. He serves as President of "
    "Stanford Angels and Entrepreneurs Texas Chapter and has deployed AI automation "
    "systems across healthcare, financial services, and industrial operations. He is "
    "also adjunct faculty at Stanford\u2019s d.school and RPI, where he teaches "
    "entrepreneurship and design thinking. He brings the builder\u2019s perspective: how to "
    "think about what to make, and how to make it fast. Learn more about Alec\u2019s work "
    "at https://inly.build/ and https://inly.education/."
)

p = doc.add_paragraph()
run = p.add_run("Alex Romine")
run.bold = True
doc.add_paragraph(
    "Alex is a cloud infrastructure and security engineer with deep experience building "
    "systems that hold up in production. He has led platform engineering, deployment "
    "automation, and security architecture at Uplight, Corteva, and Unsupervised, and "
    "works with Google as a consultant on enterprise deployments. His toolkit spans AWS, "
    "GCP, Kubernetes, and Terraform. He is the person in the room who can tell you not "
    "just how to build something, but how to build it right."
)

# --- What to Bring ---
doc.add_heading("What to Bring", level=1)
doc.add_paragraph("A laptop (8GB+ RAM, 5GB+ free storage)", style="List Bullet")
doc.add_paragraph("A Gmail account", style="List Bullet")
doc.add_paragraph("A topic, goal, or problem space you care about", style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("About cost: ")
run.bold = True
p.add_run(
    "The core tool we build together is completely free to run: no API keys, no "
    "subscriptions, no credit card required. The stretch goal (AI-powered engagement "
    "suggestions) uses a paid API at roughly $0.01 per use. If you want to try it "
    "during the workshop or keep using it afterward, budget ~$25 for an API key, but "
    "this is entirely optional. You\u2019ll walk out with a fully working tool either way."
)

# --- Pre-Work ---
doc.add_heading("Pre-Work: Before You Arrive", level=1)
doc.add_paragraph(
    "Please complete the following before the workshop. It will take about 20 minutes "
    "and will mean we spend our time building rather than setting up. We\u2019ll send "
    "everyone who registers a video walking you through the exact steps."
)

prework = [
    (
        "1. Download Claude Desktop",
        "Go to claude.ai/download and install the free desktop app. This is the AI "
        "coding partner we\u2019ll use to build the tool. Sign up for a free account using "
        "your Gmail address.",
    ),
    (
        "2. Install Python",
        "We\u2019ll send registered attendees a short setup guide before the event with "
        "step-by-step instructions for installing Python on your laptop. If you already "
        "have Python 3.10 or newer, you\u2019re set.",
    ),
    (
        "3. (Optional) Set up API access for the stretch goal",
        "If you want to try the AI-powered engagement feature during or after the "
        "workshop, you\u2019ll need an API key. We\u2019ll send setup instructions before the "
        "event. Budget ~$25 for credits. This is not required for the core tool.",
    ),
    (
        "4. Come with a topic in mind",
        "Think about a domain you want to go deeper in. It could be an industry you "
        "work in, a technology you are tracking, a health or wellness area, a hobby, a "
        "cause, or any subject where you wish you had a better handle on who the real "
        "experts are and what they are saying. The more specific your topic, the more "
        "useful the tool will be. \u201cAI\u201d is too broad. \u201cAI applied to drug discovery\u201d "
        "or \u201cearly childhood literacy research\u201d will get you something you can actually use.",
    ),
    (
        "5. Ideas for other projects?",
        "Consider this question: what is one thing in your work or life that you do "
        "repeatedly, that feels like it should not require a human every time? This could "
        "be something like sorting through emails of a certain type, summarizing documents "
        "before meetings, generating a first draft of something you write often, or "
        "pulling together information from multiple places. If you move through the main "
        "build quickly and want to keep going, that idea is your next project. A good "
        "scope for a first independent build is something that saves you 30 minutes a "
        "week, involves one or two inputs (a file, a search query, a form), and produces "
        "one clear output (a summary, a draft, a sorted list). Anything larger than that "
        "is worth breaking into smaller pieces first.",
    ),
]
for title, body in prework:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(body)

# --- FAQ ---
doc.add_heading("Frequently Asked Questions", level=1)

faqs = [
    (
        "Do I need any coding experience?",
        "None at all. That is the point. Vibe-coding lets you build with natural "
        "language. If you can describe what you want, you can build it.",
    ),
    (
        "What if I fall behind during the build?",
        "No problem. You will receive the instructor\u2019s completed version of the tool "
        "regardless of how far you get with your own build. The goal is that everyone "
        "leaves with something they can use.",
    ),
    (
        "Will I be able to keep and use the tool after the event?",
        "Yes. The tool runs on your laptop as a Python program. Both your version and "
        "the instructor\u2019s polished version are yours to keep, modify, and use however "
        "you like.",
    ),
    (
        "Is this worthwhile if I already have a technical background?",
        "Experienced developers often find these sessions valuable for different reasons: "
        "exploring AI-assisted workflows, thinking through architecture decisions, and "
        "discussing where vibe-coding fits alongside traditional engineering. Alex Romine "
        "will be on hand for the deeper technical conversations.",
    ),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

doc.save(r"C:\Users\alecm\Downloads\vibe_coding_workshop.docx")
print("Saved to C:\\Users\\alecm\\Downloads\\vibe_coding_workshop.docx")
